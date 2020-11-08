from docx import Document
import wikipedia
#wikipedia.set_lang('th')
#print(wikipedia.summary('biden'))

def Wiki(keyword,lang='th'):
    wikipedia.set_lang(lang)
    data = wikipedia.summary(keyword)

    # page + content
    data2 = wikipedia.page(keyword)
    data2 = data2.content

    doc = Document()
    doc.add_heading(keyword,0)

    doc.add_paragraph(data2)
    doc.save(keyword+'.docx')
    print('สร้างไฟล์สำเร็จ')

from tkinter import *
from tkinter import ttk
from tkinter import messagebox

GUI = Tk()
GUI.title('BunBun Wiki')
GUI.geometry('400x300')

#-----------Config-------------
FONT1 = ('Angsana New',15)

#-----------Label--------------
L = ttk.Label(GUI,text='Search Wiki',font=FONT1)
L.pack()

#-----------Search-------------
v_search = StringVar()
E1 = ttk.Entry(GUI,textvariable=v_search,font=FONT1,width=40)
E1.pack(pady=10)

#-----------Search Fn-------------
def Search():
    keyword = v_search.get()
    try: #if success let it pass
        language = v_radio.get() #th en zh
        Wiki(keyword,language)
    except: #show warning when error
        messagebox.showwarning('Keyword Error','Search again')
        
        
    #print(wikipedia.search(keyword))
    #result = wikipedia.summary(keyword,sentences=1)
    #print(result)

#-----------Button------------
B1 = ttk.Button(GUI,text='Search',command=Search)
B1.pack(ipadx=20,ipady=10)

#----------Radio Button------------
F1 = Frame(GUI)
F1.pack(pady=10)

v_radio = StringVar()
RB1 = ttk.Radiobutton(F1,text='Thai',variable=v_radio,value='th')
RB2 = ttk.Radiobutton(F1,text='Eng',variable=v_radio,value='en')
RB3 = ttk.Radiobutton(F1,text='Chinese',variable=v_radio,value='zh')
RB1.invoke() #default value to Thai
'''
RB1.pack()
RB2.pack()
RB3.pack()
'''
RB1.grid(row=0,column=0)
RB2.grid(row=0,column=1)
RB3.grid(row=0,column=2)

#-----------End-------------
GUI.mainloop()
