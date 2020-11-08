from docx import Document
import wikipedia
'''
    from docx import Document
    file: docx.py
        class Document():

    import docx
    doc = docx.Document()

    import wikipedia
    file: wikipedia.py
        def summary():
'''
def Wiki(keyword,lang='th'):
    wikipedia.set_lang(lang)
    data = wikipedia.summary(keyword)

    # page + content
    data2 = wikipedia.page(keyword)
    data2 = data2.content

    doc = Document()
    doc.add_heading(keyword,0)

    doc.add_paragraph(data2)
    doc.save(keyword+'.docx')
    print('สร้างไฟล์สำเร็จ')

Wiki('Similan islands','en')
